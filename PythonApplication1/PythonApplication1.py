import os
import json
from docx import Document
from docx.shared import Inches
from PIL import Image
from io import BytesIO

# Функция для извлечения изображений из документа и замены их на плейсхолдеры
def extract_images(doc_path, output_folder):
    doc = Document(doc_path)
    image_data = {}  # Словарь для хранения информации об изображениях
    count = 1  # Счетчик изображений

    if not os.path.exists(output_folder):
        os.makedirs(output_folder)  # Создание папки, если её нет

    # Проход по абзацам документа
    for para in doc.paragraphs:
        for run in para.runs:
            drawing = run._element.findall(".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing")
            if drawing:
                # Ищем в `drawing` тег `blip`, который указывает на ID картинки в `rels`
                blip = drawing[0].find(".//{http://schemas.openxmlformats.org/drawingml/2006/main}blip")
                if blip is not None:
                    rId = blip.get("{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed")
                    if rId in doc.part.rels:
                        image = doc.part.rels[rId].target_part.blob  # Достаем данные изображения
                        img_id = f"IMAGE_{count}"
                        img_filename = os.path.join(output_folder, f"{img_id}.png")

                        # Сохраняем изображение
                        with open(img_filename, "wb") as img_file:
                            img_file.write(image)

                        image_data[img_id] = img_filename  # Запоминаем путь
                        run.text = f"[[{img_id}]]"  # Вставляем плейсхолдер
                        count += 1  # Увеличиваем счетчик

    # Сохраняем JSON с соответствием плейсхолдеров и файлов
    with open(os.path.join(output_folder, "image_data.json"), "w") as json_file:
        json.dump(image_data, json_file, indent=4)

    # Сохраняем новый документ с плейсхолдерами
    new_doc_path = os.path.join(output_folder, "document_with_placeholders.docx")
    doc.save(new_doc_path)
    print(f"✅ Images extracted and replaced with placeholders. Saved at: {new_doc_path}")
   

# Функция для вставки изображений обратно в документ
def insert_images(doc_path, image_folder):
	try:
		# Чтение данных о изображениях из JSON файла
		with open(os.path.join(image_folder, "image_data.json"), "r") as json_file:
			image_data = json.load(json_file)
		
		# Открытие документа Word с плейсхолдерами
		doc = Document(doc_path)
		
		# Замена плейсхолдеров на изображения
		for para in doc.paragraphs:
			for img_id, img_path in image_data.items():
				if f"[[{img_id}]]" in para.text:  # Проверка на наличие плейсхолдера
					para.text = para.text.replace(f"[[{img_id}]]", "")  # Удаление плейсхолдера
					run = para.add_run()  # Добавление нового "run" (части текста)
					run.add_picture(img_path, width=Inches(2))  # Вставка изображения с указанной шириной (можно сделать адаптивным)
		
		# Сохранение документа с изображениями
		new_doc_path = os.path.join(image_folder, "document_with_images.docx")
		doc.save(new_doc_path)
		print(f"✅ Images reinserted. Saved at: {new_doc_path}")
	except Exception as e:
		print(f"❌ Error: {e}")  # Вывод ошибки в случае неудачи

# Основная функция, которая позволяет выбрать режим работы программы
def main():
	print("Choose mode:")
	print("1 - Extract images and replace with placeholders")
	print("2 - Insert images back into document")
	mode = input("Enter choice (1/2): ")
	
	if mode == "1":  # Режим извлечения изображений и замены на плейсхолдеры
		doc_path = input("Enter path to Word document: ").strip()
		output_folder = input("Enter folder to save images and new document: ").strip()
		extract_images(doc_path, output_folder)  # Вызов функции извлечения изображений
	elif mode == "2":  # Режим вставки изображений обратно в документ
		doc_path = input("Enter path to document with placeholders: ").strip()
		image_folder = input("Enter folder with saved images and JSON data: ").strip()
		insert_images(doc_path, image_folder)  # Вызов функции вставки изображений
	else:
		print("❌ Invalid choice. Exiting.")  # Некорректный выбор, выход
	
# Запуск программы
if __name__ == "__main__":
	main()
 